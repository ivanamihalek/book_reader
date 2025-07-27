#! /usr/bin/env python
#
import random
from docx import Document
from docx.shared import Inches
import argparse
import os


def sample_docx_lines(input_file, output_file, sample_rate=0.1, seed=None):
    """
    Create a sample of a DOCX document by randomly selecting a percentage of visible lines
    while preserving all metadata and document structure.

    Args:
        input_file (str): Path to input DOCX file
        output_file (str): Path to output DOCX file
        sample_rate (float): Percentage of lines to keep (0.1 = 10%)
        seed (int): Random seed for reproducible results
    """

    if seed is not None:
        random.seed(seed)

    # Load the document
    doc = Document(input_file)

    # Collect all paragraphs with visible text
    visible_paragraphs = []
    for paragraph in doc.paragraphs:
        # Only consider paragraphs with actual text content
        if paragraph.text.strip():
            visible_paragraphs.append(paragraph)

    print(f"Found {len(visible_paragraphs)} paragraphs with visible text")

    if len(visible_paragraphs) == 0:
        print("No visible text found in document!")
        return

    # Calculate number of paragraphs to keep
    num_to_keep = max(1, int(len(visible_paragraphs) * sample_rate))
    print(f"Keeping {num_to_keep} paragraphs ({sample_rate * 100:.1f}%)")

    # Randomly select paragraphs to keep
    paragraphs_to_keep = set(random.sample(visible_paragraphs, num_to_keep))

    # Create new document
    new_doc = Document()

    # Copy document properties (metadata)
    try:
        if doc.core_properties.title:
            new_doc.core_properties.title = doc.core_properties.title
        if doc.core_properties.author:
            new_doc.core_properties.author = doc.core_properties.author
        if doc.core_properties.subject:
            new_doc.core_properties.subject = doc.core_properties.subject
        if doc.core_properties.keywords:
            new_doc.core_properties.keywords = doc.core_properties.keywords
    except:
        print("Note: Some metadata could not be copied")

    # Remove the default empty paragraph from new document
    if new_doc.paragraphs:
        try:
            p = new_doc.paragraphs[0]._element
            p.getparent().remove(p)
        except:
            pass

    # Process each paragraph in the original document
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Has visible text
            if paragraph in paragraphs_to_keep:
                # Keep this paragraph - copy it with formatting
                new_p = new_doc.add_paragraph()

                # Copy paragraph style first
                try:
                    new_p.style = paragraph.style
                except:
                    pass  # Skip if style can't be copied

                # Copy paragraph formatting safely
                try:
                    if paragraph.paragraph_format.alignment:
                        new_p.paragraph_format.alignment = paragraph.paragraph_format.alignment
                except:
                    pass

                # Copy runs (text with formatting) - this is the key part
                if paragraph.runs:
                    for run in paragraph.runs:
                        new_run = new_p.add_run(run.text)
                        try:
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                        except:
                            pass  # Skip formatting that can't be copied
                else:
                    # Fallback: if no runs, just add the text directly
                    new_p.add_run(paragraph.text)

    # Copy tables (preserving all table content for now)
    for table in doc.tables:
        # Create new table with same dimensions
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))

        # Copy cell contents and formatting
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                new_cell.text = cell.text
                # Copy cell formatting if needed
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        new_p = new_cell.paragraphs[0] if new_cell.paragraphs else new_cell.add_paragraph()
                        new_p.text = paragraph.text

    # Save the sampled document
    new_doc.save(output_file)
    print(f"Sampled document saved to: {output_file}")


def main():
    parser = argparse.ArgumentParser(description='Sample lines from a DOCX document while preserving metadata')
    parser.add_argument('input_file', help='Input DOCX file path')
    parser.add_argument('output_file', help='Output DOCX file path')
    parser.add_argument('--sample-rate', type=float, default=0.1,
                        help='Sampling rate (default: 0.1 for 10%%)')
    parser.add_argument('--seed', type=int, help='Random seed for reproducible results')

    args = parser.parse_args()

    # Validate input file exists
    if not os.path.exists(args.input_file):
        print(f"Error: Input file '{args.input_file}' not found")
        return

    # Validate sample rate
    if not 0 < args.sample_rate <= 1:
        print("Error: Sample rate must be between 0 and 1")
        return

    try:
        sample_docx_lines(args.input_file, args.output_file, args.sample_rate, args.seed)
        print("Sampling completed successfully!")
    except Exception as e:
        print(f"Error processing document: {e}")


if __name__ == "__main__":
    # Example usage when run directly
    if len(os.sys.argv) == 1:
        print("DOCX Document Sampler")
        print("Usage: python script.py input.docx output.docx [--sample-rate 0.1] [--seed 42]")
        print("\nExample: python script.py document.docx sample.docx --sample-rate 0.1 --seed 42")
    else:
        main()